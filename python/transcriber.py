import numpy as np
import torch
import whisper
import io
from docx import Document
from typing import Dict, Any

class Transcriber:
    def __init__(self):
        self.model = whisper.load_model("base")
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"Using device: {self.device}")

    def transcribe_audio(self, audio_data: bytes) -> Dict[str, Any]:
        try:
            # Convert audio bytes to numpy array
            audio_np = np.frombuffer(audio_data, dtype=np.float32)
            
            # Perform transcription
            result = self.model.transcribe(audio_np)
            
            return {
                "text": result["text"],
                "segments": result["segments"],
                "language": result["language"]
            }
        except Exception as e:
            print(f"Transcription error: {str(e)}")
            return {"error": str(e)}

    def export_to_docx(self, transcription: Dict[str, Any]) -> bytes:
        doc = Document()
        doc.add_heading('Transcription', 0)
        
        # Add main transcription text
        doc.add_paragraph(transcription["text"])
        
        # Add segments with timestamps if available
        if "segments" in transcription:
            doc.add_heading('Segments', level=1)
            for segment in transcription["segments"]:
                start = format_timestamp(segment["start"])
                end = format_timestamp(segment["end"])
                text = segment["text"]
                doc.add_paragraph(f"[{start} - {end}] {text}")
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        return doc_bytes.getvalue()

    def export_to_txt(self, transcription: Dict[str, Any]) -> str:
        output = ["Transcription:", transcription["text"], "", "Segments:"]
        
        if "segments" in transcription:
            for segment in transcription["segments"]:
                start = format_timestamp(segment["start"])
                end = format_timestamp(segment["end"])
                text = segment["text"]
                output.append(f"[{start} - {end}] {text}")
        
        return "\n".join(output)

def format_timestamp(seconds: float) -> str:
    minutes = int(seconds // 60)
    seconds = int(seconds % 60)
    return f"{minutes:02d}:{seconds:02d}"

# Initialize transcriber
transcriber = Transcriber()